from pathlib import Path

import docx

from src.pipeline.ngs_parser import parse_ngs_docx


def _make_docx(tmp_path: Path) -> Path:
    doc = docx.Document()
    doc.add_paragraph("1AA Health exam", style="Heading 1")
    doc.add_paragraph("Routine periodic visit for an adult patient.")
    doc.add_paragraph("Includes codes: K040, K005")
    doc.add_paragraph("1AB Consultation", style="Heading 1")
    doc.add_paragraph("Specialist consultation.")
    doc.add_paragraph("Includes codes: A005, 01712")
    out = tmp_path / "ngs.docx"
    doc.save(str(out))
    return out


def test_parse_ngs_docx_emits_records(tmp_path):
    path = _make_docx(tmp_path)
    records = parse_ngs_docx(path)
    assert len(records) == 2
    assert records[0].ngs_code == "1AA"
    assert records[0].ngs_label == "Health exam"
    assert "K040" in records[0].code_refs
    assert "01712" in records[1].code_refs
