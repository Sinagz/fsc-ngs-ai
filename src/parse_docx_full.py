"""
Parse both DOCX reference files, including ALL tables.

Outputs:
  data/parsed/ngs/ngs_categories.json       – NGS code definitions
  data/parsed/ngs/ccp_ngs_map.json          – CCP modifier → NGS category per province
  data/parsed/ngs/province_rules.json       – FSC structure + role-code rules per province
  data/parsed/ngs/role_code_tables.json     – role code tables per province
"""
import json
import re
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
DOCX_DIR = BASE_DIR / "data" / "raw" / "docx"
OUT_DIR  = BASE_DIR / "data" / "parsed" / "ngs"
OUT_DIR.mkdir(parents=True, exist_ok=True)

GROUPING_RULES_DOCX  = DOCX_DIR / "Fee Codes and Grouping Rules.docx"
NGS_CATEGORIES_DOCX  = DOCX_DIR / "NPDB National Grouping System Categories.docx"


def norm(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u2013", "-").replace("\u2019", "'")
    text = text.replace("\u2014", "-").replace("\u00e9", "e").replace("\ufffd", "")
    return re.sub(r"\s+", " ", text).strip()


# ── NGS categories ─────────────────────────────────────────────────────────

NGS_CODE_LINE = re.compile(r"^(\d{3}\.\d{3})\s+(.+)$")

def parse_ngs_categories(path: Path) -> list:
    doc = Document(str(path))
    rows = []
    current_section = ""
    current_section_desc = ""
    current_cat = None

    for para in doc.paragraphs:
        text = norm(para.text)
        if not text:
            continue

        # Heading styles mark sections
        style = para.style.name
        if style.startswith("Heading"):
            if current_cat:
                rows.append(current_cat)
                current_cat = None
            current_section = text
            current_section_desc = ""
            continue

        m = NGS_CODE_LINE.match(text)
        if m:
            if current_cat:
                rows.append(current_cat)
            current_cat = {
                "ngs_code":  m.group(1),
                "ngs_label": norm(m.group(2)),
                "ngs_section": current_section,
                "ngs_description": "",
            }
        else:
            # Descriptive line: attach to current category or section
            if current_cat:
                if not current_cat["ngs_description"]:
                    current_cat["ngs_description"] = text
                else:
                    current_cat["ngs_description"] += " " + text
            else:
                current_section_desc = (current_section_desc + " " + text).strip()

    if current_cat:
        rows.append(current_cat)

    # Deduplicate by ngs_code
    seen = {}
    out  = []
    for r in rows:
        if r["ngs_code"] not in seen:
            seen[r["ngs_code"]] = True
            out.append(r)
    return out


# ── Grouping rules (paragraphs + tables) ───────────────────────────────────

PROVINCE_HEADINGS = {
    "Newfoundland": "NL", "Prince Edward Island": "PE", "Nova Scotia": "NS",
    "New Brunswick": "NB", "Quebec": "QC", "Qu\u00e9bec": "QC",
    "Ontario": "ON", "Manitoba": "MB", "Saskatchewan": "SK",
    "Alberta": "AB", "British Columbia": "BC", "Yukon": "YT",
}

CCP_NGS_RE = re.compile(r"^NGS\s*(\d{3}\.\d{3})", re.IGNORECASE)
NGS_INLINE  = re.compile(r"NGS\s+(\d{3}\.\d{3})", re.IGNORECASE)


def _table_to_dicts(table) -> list[dict]:
    """Convert a docx table to a list of row dicts using the first row as header."""
    rows = []
    headers = []
    for i, row in enumerate(table.rows):
        cells = [norm(c.text) for c in row.cells]
        # merge adjacent duplicate cells (caused by merged cells in docx)
        cells = [cells[j] for j in range(len(cells))
                 if j == 0 or cells[j] != cells[j-1]]
        if i == 0:
            headers = cells
        else:
            if any(c for c in cells):
                rows.append(dict(zip(headers, cells)))
    return rows


def _extract_ngs_from_text(text: str) -> list[str]:
    return NGS_INLINE.findall(text)


def parse_grouping_rules(path: Path):
    doc = Document(str(path))

    ccp_ngs_maps   = []   # {province, ccp, description, ngs_codes, notes}
    province_rules = []   # {province, rule_type, detail}
    role_tables    = []   # {province, role_code, description, ngs_codes, notes}

    # Walk body elements in order (paragraphs and tables interleaved)
    current_province = "ALL"
    current_section  = ""

    for elem in doc.element.body:
        tag = elem.tag.split("}")[-1]

        # ── paragraph ──
        if tag == "p":
            from docx.oxml.ns import qn
            from docx.text.paragraph import Paragraph
            para = Paragraph(elem, doc)
            text = norm(para.text)
            if not text:
                continue

            style = para.style.name

            # Province headings
            if text in PROVINCE_HEADINGS:
                current_province = PROVINCE_HEADINGS[text]
                continue
            if style.startswith("Heading"):
                current_section = text
                if text in PROVINCE_HEADINGS:
                    current_province = PROVINCE_HEADINGS[text]
                continue

            # Global role-code rules
            if "surgical assistance" in text.lower():
                ngs = _extract_ngs_from_text(text)
                province_rules.append({
                    "province": "ALL",
                    "rule_type": "role_code",
                    "trigger":   "surgical assistance role code",
                    "ngs_codes": ngs,
                    "detail":    text,
                })
            if "anesthesia" in text.lower() or "anaesthesia" in text.lower():
                ngs = _extract_ngs_from_text(text)
                if ngs:
                    province_rules.append({
                        "province": "ALL",
                        "rule_type": "role_code",
                        "trigger":   "anesthesia role code",
                        "ngs_codes": ngs,
                        "detail":    text,
                    })

            # FSC structure rules
            if any(kw in text.lower() for kw in
                   ["fsc code", "fee code", "leading zero", "role code",
                    "first four", "first three", "fifth digit", "padded"]):
                province_rules.append({
                    "province": current_province,
                    "rule_type": "fsc_structure",
                    "trigger":   text[:80],
                    "ngs_codes": [],
                    "detail":    text,
                })

        # ── table ──
        elif tag == "tbl":
            from docx.table import Table
            table = Table(elem, doc)
            rows  = _table_to_dicts(table)
            if not rows:
                continue

            headers_lower = {k.lower() for k in rows[0].keys()}

            # CCP modifier table
            if any("ccp" in h or "modifier" in h for h in headers_lower):
                for row in rows:
                    ccp  = row.get("CCP Modifier") or row.get("CCP modifier") or ""
                    desc = (row.get("Description") or row.get("Provincial Modifier and Notes")
                            or row.get("Provincial Modifier") or "")
                    ngs_raw = (row.get("NGS Category") or row.get("NGS Category.1") or "")
                    ngs_codes = re.findall(r"\d{3}\.\d{3}", ngs_raw)
                    if ccp and (ngs_codes or desc):
                        ccp_ngs_maps.append({
                            "province":  current_province,
                            "ccp_code":  norm(ccp),
                            "ccp_desc":  norm(desc),
                            "ngs_codes": ngs_codes,
                            "raw_ngs":   norm(ngs_raw),
                        })

            # Role code table
            elif any("role" in h for h in headers_lower):
                for row in rows:
                    rc   = (row.get("Role Code") or "")
                    rdesc= (row.get("Description") or "")
                    rnotes=(row.get("Notes") or "")
                    ngs_codes = re.findall(r"\d{3}\.\d{3}", rnotes + " " + rdesc)
                    if rc:
                        role_tables.append({
                            "province":   current_province,
                            "role_code":  norm(rc),
                            "description":norm(rdesc),
                            "ngs_codes":  ngs_codes,
                            "notes":      norm(rnotes),
                        })

            # FSC example table
            elif any("fsc" in h.lower() for h in headers_lower):
                for row in rows:
                    fsc  = (row.get("FSC Code") or "")
                    fee  = (row.get("Fee Code") or "")
                    rc   = (row.get("Role Code") or "")
                    if fsc or fee:
                        province_rules.append({
                            "province": current_province,
                            "rule_type": "fsc_example",
                            "trigger":   f"FSC={fsc} => fee={fee} role={rc}",
                            "ngs_codes": [],
                            "detail":    str(row),
                        })

    return ccp_ngs_maps, province_rules, role_tables


# ── main ───────────────────────────────────────────────────────────────────

def main():
    # 1. NGS categories
    print("Parsing NGS categories...", flush=True)
    ngs_cats = parse_ngs_categories(NGS_CATEGORIES_DOCX)
    out = OUT_DIR / "ngs_categories.json"
    out.write_text(json.dumps(ngs_cats, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  {len(ngs_cats)} NGS categories -> {out}")

    # 2. Grouping rules
    print("Parsing grouping rules + tables...", flush=True)
    ccp_maps, prov_rules, role_tables = parse_grouping_rules(GROUPING_RULES_DOCX)

    out2 = OUT_DIR / "ccp_ngs_map.json"
    out2.write_text(json.dumps(ccp_maps, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  {len(ccp_maps)} CCP-NGS map entries -> {out2}")

    out3 = OUT_DIR / "province_rules.json"
    out3.write_text(json.dumps(prov_rules, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  {len(prov_rules)} province rules -> {out3}")

    out4 = OUT_DIR / "role_code_tables.json"
    out4.write_text(json.dumps(role_tables, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"  {len(role_tables)} role code table entries -> {out4}")


if __name__ == "__main__":
    main()
